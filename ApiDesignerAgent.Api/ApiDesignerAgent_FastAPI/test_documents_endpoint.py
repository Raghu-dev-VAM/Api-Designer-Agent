import io, json
from fastapi.testclient import TestClient
from main import app

client = TestClient(app, raise_server_exceptions=True)
CSV = b"Story ID,User Story,Priority,Acceptance Criteria\nUS-001,As a user I want to login,High,Can login with valid creds\nUS-002,As a user I want to register,Medium,Can create account\n"

results = []
def check(label, cond, detail=""):
    results.append(("[PASS] " if cond else "[FAIL] ") + label + (f" | {detail}" if detail else ""))

# A: CSV no mapping → needs_mapping
r = client.post("/api/documents/extract", files={"file": ("s.csv", CSV, "text/csv")})
b = r.json()
check("A 200",              r.status_code == 200,               f"got {r.status_code}")
check("A needs_mapping",    b.get("needs_mapping") == True,     str(b.get("needs_mapping")))
check("A columns list",     "User Story" in b.get("columns",[]),str(b.get("columns")))

# B: CSV full mapping → 2 requirements
m = json.dumps({"userStory":"User Story","storyId":"Story ID","title":"","priority":"Priority","acceptanceCriteria":"Acceptance Criteria"})
r = client.post("/api/documents/extract", files={"file":("s.csv",CSV,"text/csv")}, data={"columnMapping":m})
b = r.json()
check("B 200",              r.status_code == 200,               f"got {r.status_code}")
check("B 2 reqs",           len(b.get("requirements",[])) == 2, f"got {len(b.get('requirements',[]))}")
if b.get("requirements"):
    req = b["requirements"][0]
    check("B id=US-001",    req["id"] == "US-001",   req.get("id"))
    check("B priority=High",req["priority"]=="High", req.get("priority"))
    check("B desc filled",  bool(req.get("desc")),   req.get("desc","")[:60])

# C: userStory-only mapping (all others blank)
m_min = json.dumps({"userStory":"User Story","storyId":"","title":"","priority":"","acceptanceCriteria":""})
r = client.post("/api/documents/extract", files={"file":("s.csv",CSV,"text/csv")}, data={"columnMapping":m_min})
b = r.json()
check("C 200",              r.status_code == 200,               f"got {r.status_code}")
check("C 2 reqs",           len(b.get("requirements",[])) == 2, f"got {len(b.get('requirements',[]))}")

# D: unsupported ext → 400
r = client.post("/api/documents/extract", files={"file":("doc.txt",b"hello","text/plain")})
check("D 400 bad ext",      r.status_code == 400,               f"got {r.status_code}")

# E: empty userStory in mapping → 400
bad = json.dumps({"userStory":"","storyId":"","title":"","priority":"","acceptanceCriteria":""})
r = client.post("/api/documents/extract", files={"file":("s.csv",CSV,"text/csv")}, data={"columnMapping":bad})
check("E 400 no userStory", r.status_code == 400,               f"got {r.status_code}")

# F: XLSX no mapping → needs_mapping
import openpyxl
wb = openpyxl.Workbook(); ws = wb.active
ws.append(["Story ID","User Story","Priority"])
ws.append(["US-010","As a user I want to reset password","High"])
buf = io.BytesIO(); wb.save(buf); xb = buf.getvalue()
r = client.post("/api/documents/extract", files={"file":("s.xlsx",xb,"application/octet-stream")})
b = r.json()
check("F xlsx needs_mapping", b.get("needs_mapping")==True, str(b.get("needs_mapping")))

# G: XLSX with mapping → 1 requirement
m2 = json.dumps({"userStory":"User Story","storyId":"Story ID","title":"","priority":"Priority","acceptanceCriteria":""})
r = client.post("/api/documents/extract", files={"file":("s.xlsx",xb,"application/octet-stream")}, data={"columnMapping":m2})
b = r.json()
check("G xlsx 200",         r.status_code == 200,               f"got {r.status_code}")
check("G xlsx 1 req",       len(b.get("requirements",[])) == 1, f"got {len(b.get('requirements',[]))}")

# H: DOCX → extracts requirements (calls Groq, so just check 200 + non-empty)
import docx as _docx
doc = _docx.Document()
doc.add_paragraph("As a user I want to login so that I can access my account")
doc.add_paragraph("As an admin I want to manage users so that I can control access")
buf2 = io.BytesIO(); doc.save(buf2); db = buf2.getvalue()
r = client.post("/api/documents/extract", files={"file":("t.docx",db,"application/octet-stream")})
b = r.json()
check("H docx 200",         r.status_code == 200,               f"got {r.status_code}")
check("H docx has reqs",    len(b.get("requirements",[])) > 0,  f"got {len(b.get('requirements',[]))}")
check("H docx extraction_method", b.get("extraction_method")=="word", b.get("extraction_method"))

print("\n=== RESULTS ===")
for line in results: print(line)
failed = [l for l in results if l.startswith("[FAIL]")]
print(f"\n{len(results)-len(failed)}/{len(results)} passed")
if failed:
    print("FAILURES:")
    for f in failed: print(" ", f)
