import csv
import io
import json
import logging
import re
from datetime import datetime, timezone
from typing import Optional

import openpyxl
import yaml as pyyaml
from fastapi import APIRouter, File, Form, HTTPException, Request, UploadFile

from models import (
    GenerateRequest, GenerateResponse, Requirement,
    ValidateRequest, ValidateResponse,
    ArtifactRequest,
)
from dependencies import get_groq_service, get_python_service

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/designer", tags=["designer"])


def _get_groq():
    return get_groq_service()


def _get_python():
    return get_python_service()


@router.post("/generate", response_model=GenerateResponse)
async def generate_openapi(request: GenerateRequest):
    # Resolve requirements from raw_input if provided
    if request.raw_input and request.raw_input.strip():
        request = _inject_raw_input(request)

    if not request.requirements:
        raise HTTPException(
            status_code=400,
            detail="Provide either 'requirements' array or 'raw_input' (JSON / YAML / plain text)."
        )

    # Normalise output_format
    # None / omitted  → both
    # 'yaml'          → YAML only
    # 'json'          → JSON only
    fmt = (request.output_format or "").strip().lower() or "both"
    if fmt not in ("yaml", "json", "both"):
        raise HTTPException(status_code=400, detail="output_format must be 'yaml', 'json', or omitted (both).")

    to_process = [r for r in request.requirements if (r.status or "Draft").lower() != "rejected"]
    if not to_process:
        raise HTTPException(status_code=400, detail="All requirements are rejected.")

    logger.info("Generating OpenAPI (format=%s) for %d requirement(s)", fmt, len(to_process))

    try:
        yaml_raw   = await _get_groq().generate_openapi(
            GenerateRequest(requirements=to_process,
                            api_title=request.api_title,
                            api_version=request.api_version)
        )
        yaml_clean = _clean_yaml(yaml_raw)
        summary    = await _get_groq().generate_summary(yaml_clean)

        # Build response based on requested format
        if fmt == "yaml":
            return GenerateResponse(
                open_api_yaml=yaml_clean,
                open_api_json=None,
                output_format="yaml",
                summary=summary,
                generated_at=datetime.now(timezone.utc).isoformat(),
            )
        elif fmt == "json":
            json_spec = _get_python().convert_yaml_to_json(yaml_clean)
            return GenerateResponse(
                open_api_yaml=None,
                open_api_json=json_spec,
                output_format="json",
                summary=summary,
                generated_at=datetime.now(timezone.utc).isoformat(),
            )
        else:  # both
            json_spec = _get_python().convert_yaml_to_json(yaml_clean)
            return GenerateResponse(
                open_api_yaml=yaml_clean,
                open_api_json=json_spec,
                output_format="both",
                summary=summary,
                generated_at=datetime.now(timezone.utc).isoformat(),
            )

    except Exception as ex:
        logger.error("Generate failed: %s", ex)
        msg = str(ex)
        if "Invalid API Key" in msg or "invalid_api_key" in msg:
            raise HTTPException(status_code=503, detail="LLM service unavailable: invalid or expired Groq API key.")
        if "All LLM providers" in msg:
            raise HTTPException(status_code=503, detail="LLM service temporarily unavailable. Please try again.")
        raise HTTPException(status_code=500, detail=msg)


# ── raw_input helpers ─────────────────────────────────────────────────────────

def _inject_raw_input(request: GenerateRequest) -> GenerateRequest:
    """Parse raw_input (JSON / YAML / plain text) into Requirement objects."""
    raw  = request.raw_input.strip()
    reqs: list[Requirement] = []

    # Try JSON
    try:
        parsed = json.loads(raw)
        if isinstance(parsed, list):
            reqs = [_dict_to_req(i, d) for i, d in enumerate(parsed)]
        elif isinstance(parsed, dict):
            reqs = [_dict_to_req(0, parsed)]
    except (json.JSONDecodeError, ValueError):
        pass

    # Try YAML
    if not reqs:
        try:
            parsed = pyyaml.safe_load(raw)
            if isinstance(parsed, list):
                reqs = [_dict_to_req(i, d) for i, d in enumerate(parsed) if isinstance(d, dict)]
            elif isinstance(parsed, dict):
                reqs = [_dict_to_req(0, parsed)]
        except Exception:
            pass

    # Plain text fallback
    if not reqs:
        reqs = [Requirement(
            id="R-001", title=request.api_title or "Requirement",
            description=raw[:2000], source="raw_input",
            priority="Medium", status="Draft",
        )]

    return GenerateRequest(
        requirements=reqs,
        api_title=request.api_title,
        api_version=request.api_version,
        output_format=request.output_format,
    )


def _dict_to_req(index: int, d: dict) -> Requirement:
    seq = f"{index + 1:03d}"
    return Requirement(
        id=str(d.get("id") or d.get("storyId") or f"R-{seq}"),
        title=str(d.get("title") or d.get("name") or d.get("summary") or f"Requirement {seq}"),
        description=str(d.get("description") or d.get("desc") or d.get("userStory") or ""),
        source=str(d.get("source") or "raw_input"),
        priority=str(d.get("priority") or "Medium"),
        status=str(d.get("status") or "Draft"),
    )


@router.post("/validate", response_model=ValidateResponse)
async def validate_openapi_spec(request: ValidateRequest):
    if not request.open_api_yaml or not request.open_api_yaml.strip():
        raise HTTPException(status_code=400, detail="OpenApiYaml is required.")
    return _get_python().validate_openapi(request.open_api_yaml)


@router.post("/artifact")
async def get_artifact(request: ArtifactRequest):
    if not request.open_api_yaml or not request.open_api_yaml.strip():
        raise HTTPException(status_code=400, detail="OpenApiYaml is required.")

    artifact_type = (request.artifact_type or "").lower()
    try:
        if artifact_type == "yaml":
            return {"content": request.open_api_yaml, "file_name": "openapi.yaml", "content_type": "application/x-yaml"}
        elif artifact_type == "json":
            return {"content": _get_python().convert_yaml_to_json(_clean_yaml(request.open_api_yaml)), "file_name": "openapi.json", "content_type": "application/json"}
        elif artifact_type == "postman":
            return {"content": _get_python().generate_postman_collection(request.open_api_yaml, request.api_title or "API Collection"), "file_name": "postman_collection.json", "content_type": "application/json"}
        else:
            raise HTTPException(status_code=400, detail=f"Unknown artifact type: {request.artifact_type}. Supported: yaml, json, postman")
    except HTTPException:
        raise
    except Exception as ex:
        logger.error("Artifact generation failed: %s", ex)
        raise HTTPException(status_code=500, detail=str(ex))


@router.post("/extract-requirements")
async def extract_requirements_from_document(request: Request, file: Optional[UploadFile] = File(None)):
    # Fall back to scanning all form fields for the first UploadFile
    form = await request.form()
    if file is None or not file.filename:
        for field_value in form.values():
            if isinstance(field_value, UploadFile) and field_value.filename:
                file = field_value
                break
    if file is None or not file.filename:
        raise HTTPException(status_code=400, detail="No file uploaded. Please upload a .docx or .xlsx file.")

    fname = file.filename.lower()
    is_excel = fname.endswith(".xlsx") or fname.endswith(".xls")
    is_docx = fname.endswith(".docx")
    if not is_excel and not is_docx:
        raise HTTPException(status_code=400, detail="Only .docx, .xlsx, and .xls files are supported.")

    contents = await file.read()

    # ── Excel path ────────────────────────────────────────────────────────────
    if is_excel:
        # Column mapping: caller may pass JSON like
        # {"storyId":"ID","title":"Story Title","userStory":"Description",
        #  "priority":"Priority","acceptanceCriteria":"AC"}
        mapping_raw = form.get("columnMapping") or form.get("column_mapping") or "{}"
        try:
            col_map: dict = json.loads(mapping_raw) if isinstance(mapping_raw, str) else {}
        except Exception:
            col_map = {}

        # userStory mapping is required for Excel uploads
        if not col_map.get("userStory"):
            raise HTTPException(
                status_code=400,
                detail="columnMapping.userStory is required for Excel uploads. "
                       "Provide the exact Excel column header that contains the user story text."
            )

        try:
            import openpyxl
        except ImportError:
            raise HTTPException(status_code=500, detail="openpyxl is not installed on the server.")

        try:
            with io.BytesIO(contents) as buf:
                wb = openpyxl.load_workbook(buf, read_only=True, data_only=True)
                ws = wb.active
                rows = list(ws.iter_rows(values_only=True))
                wb.close()
        except Exception as ex:
            raise HTTPException(status_code=400, detail=f"Failed to read Excel file: {ex}")

        if not rows:
            raise HTTPException(status_code=400, detail="Excel file appears to be empty.")

        headers = [str(h).strip() if h is not None else "" for h in rows[0]]

        def _find_col(candidates: list[str]) -> int:
            for name in candidates:
                mapped = col_map.get(name, "")
                if mapped and mapped in headers:
                    return headers.index(mapped)
            for name in candidates:
                for i, h in enumerate(headers):
                    if h.lower() == name.lower():
                        return i
            return -1

        idx = {
            "storyId":            _find_col(["storyId", "id", "story id", "story_id"]),
            "title":              _find_col(["title", "story title", "name"]),
            "userStory":          _find_col(["userStory", "user story", "description", "desc"]),
            "priority":           _find_col(["priority"]),
            "acceptanceCriteria": _find_col(["acceptanceCriteria", "acceptance criteria", "ac", "criteria"]),
        }

        # Validate the required userStory column resolved to an actual header
        if idx["userStory"] == -1:
            raise HTTPException(
                status_code=400,
                detail=f"Column '{col_map['userStory']}' specified in columnMapping.userStory "
                       f"was not found in the Excel headers: {headers}"
            )

        def _cell(row: tuple, key: str) -> str:
            i = idx[key]
            return str(row[i]).strip() if i >= 0 and i < len(row) and row[i] is not None else ""

        requirements = []
        for row_num, row in enumerate(rows[1:], start=1):
            user_story = _cell(row, "userStory")
            if not user_story:
                continue  # skip rows with no user story value
            title = _cell(row, "title")
            requirements.append({
                "id":                 _cell(row, "storyId") or f"FR-{row_num:03d}",
                "title":              title,
                "desc":               user_story,
                "source":             "Uploaded Excel",
                "priority":           _cell(row, "priority") or "Medium",
                "status":             "Draft",
                "method":             "get",
                "path":               "/resource",
                "summary":            title,
                "acceptanceCriteria": _cell(row, "acceptanceCriteria"),
            })

        return {"requirements": requirements, "extraction_method": "excel", "column_indices": idx}

    # ── Word (.docx) path ─────────────────────────────────────────────────────
    try:
        import docx
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx is not installed on the server.")

    try:
        with io.BytesIO(contents) as buf:
            doc = docx.Document(buf)
            parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        parts.append(row_text)
            text = "\n".join(parts)
    except Exception as ex:
        raise HTTPException(status_code=400, detail=f"Failed to read document: {ex}")

    if not text.strip():
        raise HTTPException(status_code=400, detail="Document appears to be empty.")

    prompt = f"""You are a business analyst. Extract structured functional requirements from the following document text.

IMPORTANT: Look carefully for acceptance criteria in any format:
- "Given / When / Then" style
- Numbered or bulleted lists under headings like "Acceptance Criteria", "AC:", "Done when:"
- Table rows with criteria
- Any conditions or rules that define when a story is complete

Return a JSON array only (no markdown, no explanation) where each item has:
- id: string like "FR-001"
- title: short title for the user story
- desc: full user story text ("As a... I want... So that...")
- source: "Uploaded Document"
- priority: "High", "Medium", or "Low"
- status: "Draft"
- method: appropriate HTTP method (get, post, put, patch, delete)
- path: REST API path like /resources/{{id}}
- summary: one-line API operation summary
- acceptanceCriteria: array of acceptance criteria strings — MUST be populated if any criteria exist in the document. Do NOT return an empty array if criteria are present.

Document text:
{text[:8000]}
"""
    try:
        raw = await _get_groq()._call_groq(prompt)
        cleaned = _clean_json(raw)
        requirements = json.loads(cleaned)
        for r in requirements:
            if "desc" not in r and "description" in r:
                r["desc"] = r.pop("description")
            r.setdefault("desc", "")
            r.setdefault("summary", "")
            r.setdefault("method", "get")
            r.setdefault("path", "/resource")
            r.setdefault("status", "Draft")
            ac = r.get("acceptanceCriteria", [])
            if isinstance(ac, list):
                r["acceptanceCriteria"] = " ".join(c.strip() for c in ac if c.strip())
            elif not isinstance(ac, str):
                r["acceptanceCriteria"] = ""
        return {"requirements": requirements, "raw_text": text, "extraction_method": "llm"}
    except Exception as ex:
        logger.warning("LLM extraction failed (%s), falling back to rule-based extraction", ex)
        requirements = _rule_based_extract(text)
        return {"requirements": requirements, "raw_text": text, "extraction_method": "rule-based"}


def _rule_based_extract(text: str) -> list:
    """Extract requirements from document text using simple heuristics."""
    import re

    METHOD_KEYWORDS = {
        "create": "post", "add": "post", "register": "post", "submit": "post", "upload": "post",
        "update": "put", "edit": "put", "modify": "patch", "change": "patch",
        "delete": "delete", "remove": "delete",
        "get": "get", "list": "get", "view": "get", "retrieve": "get", "fetch": "get", "search": "get",
    }
    PRIORITY_KEYWORDS = {
        "critical": "High", "must": "High", "required": "High", "mandatory": "High",
        "should": "Medium", "recommended": "Medium",
        "optional": "Low", "nice to have": "Low", "could": "Low",
    }

    # Split into sentences / bullet lines
    lines = [l.strip() for l in re.split(r'[\n•\-–]', text) if len(l.strip()) > 20]
    # Also consider numbered items like "1." or "FR-001"
    sentences = []
    for line in lines:
        for sent in re.split(r'(?<=[.!?])\s+', line):
            if len(sent.strip()) > 20:
                sentences.append(sent.strip())

    requirements = []
    seen = set()
    counter = 1

    for sentence in sentences:
        low = sentence.lower()
        # Skip headings / very short lines
        if len(sentence) < 25 or sentence.isupper():
            continue
        if sentence in seen:
            continue
        seen.add(sentence)

        # Determine HTTP method
        method = "get"
        for kw, m in METHOD_KEYWORDS.items():
            if kw in low:
                method = m
                break

        # Determine priority
        priority = "Medium"
        for kw, p in PRIORITY_KEYWORDS.items():
            if kw in low:
                priority = p
                break

        # Build a slug for the path
        words = re.findall(r'[a-zA-Z]+', sentence)
        nouns = [w.lower() for w in words if len(w) > 3 and w.lower() not in
                 {"shall", "should", "must", "will", "have", "with", "that", "this",
                  "from", "into", "able", "user", "users", "system", "allow", "able"}]
        resource = nouns[0] if nouns else "resource"
        path = f"/{resource}s" if method in ("get", "post") else f"/{resource}s/{{id}}"

        title_words = words[:6]
        title = " ".join(title_words).title()

        requirements.append({
            "id": f"FR-{counter:03d}",
            "title": title,
            "desc": sentence,
            "source": "Uploaded Document",
            "priority": priority,
            "status": "Draft",
            "method": method,
            "path": path,
            "summary": f"{method.upper()} {path} — {title}",
            "acceptanceCriteria": "",
        })
        counter += 1
        if counter > 50:  # cap at 50 requirements
            break

    return requirements


def _clean_json(raw: str) -> str:
    import re
    t = raw.strip()
    t = re.sub(r'^```[a-zA-Z]*[ \t]*\r?\n?', '', t)
    t = re.sub(r'\r?\n```[a-zA-Z]*[ \t]*(?=\r?\n|$)', '', t)
    t = re.sub(r'^```[a-zA-Z]*[ \t]*$', '', t, flags=re.MULTILINE)
    return t.strip()


def _clean_yaml(raw: str) -> str:
    import re
    t = raw.strip()
    t = re.sub(r'^```[a-zA-Z]*[ \t]*\r?\n?', '', t)
    t = re.sub(r'\r?\n```[a-zA-Z]*[ \t]*(?=\r?\n|$)', '', t)
    t = re.sub(r'^```[a-zA-Z]*[ \t]*$', '', t, flags=re.MULTILINE)
    return t.strip()


# ── /extract helpers ─────────────────────────────────────────────────────────

_WORD_EXTS  = {".docx"}
_SHEET_EXTS = {".xlsx", ".csv"}
_ALL_EXTS   = _WORD_EXTS | _SHEET_EXTS

_COL_PATTERNS: dict[str, list[str]] = {
    "userStory":          ["user story", "userstory", "user_story", "description",
                           "desc", "story", "requirement", "as a", "as an"],
    "storyId":            ["story id", "storyid", "story_id", "id", "ticket", "issue"],
    "title":              ["title", "name", "summary", "epic"],
    "priority":           ["priority", "severity", "importance"],
    "acceptanceCriteria": ["acceptance criteria", "acceptance_criteria",
                           "acceptancecriteria", "ac", "criteria", "done when",
                           "definition of done", "dod"],
}


def _parse_sheet(content: bytes, filename: str) -> list[dict]:
    ext = ("." + filename.rsplit(".", 1)[-1]).lower() if "." in filename else ""
    if ext == ".csv":
        text = content.decode("utf-8-sig", errors="replace")
        return [dict(row) for row in csv.DictReader(io.StringIO(text))]
    if ext == ".xlsx":
        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        ws = wb.active
        rows = list(ws.iter_rows(values_only=True))
        wb.close()
        if not rows:
            return []
        headers = [str(h).strip() if h is not None else f"col_{i}" for i, h in enumerate(rows[0])]
        return [
            {headers[j]: (str(cell).strip() if cell is not None else "") for j, cell in enumerate(row)}
            for row in rows[1:]
        ]
    raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}")


def _resolve_column(value: str, columns: list[str]) -> str:
    if not value:
        return ""
    if value in columns:
        return value
    vl = value.strip().lower()
    for col in columns:
        if col.strip().lower() == vl:
            return col
    for col in columns:
        cl = col.strip().lower()
        if vl in cl or cl in vl:
            return col
    return ""


def _auto_detect(columns: list[str]) -> dict[str, str]:
    result: dict[str, str] = {}
    for field, patterns in _COL_PATTERNS.items():
        for col in columns:
            if any(p in col.strip().lower() for p in patterns):
                result[field] = col
                break
    return result


def _build_mapping(raw_input: str, columns: list[str]) -> dict[str, str]:
    raw = raw_input.strip()
    if raw.startswith("{"):
        try:
            supplied: dict = json.loads(raw)
        except Exception:
            raise HTTPException(status_code=400, detail="columnMapping must be valid JSON when sent as an object.")
        mapping = {k: _resolve_column(v, columns) for k, v in supplied.items() if v}
    else:
        resolved = _resolve_column(raw, columns)
        mapping = {"userStory": resolved} if resolved else {}
    auto = _auto_detect(columns)
    for field in _COL_PATTERNS:
        if not mapping.get(field):
            mapping[field] = auto.get(field, "")
    return mapping


def _priority_map(raw: str) -> str:
    r = raw.strip().lower()
    if r in ("high", "1", "critical", "blocker"):
        return "High"
    if r in ("low", "3", "4", "minor", "trivial"):
        return "Low"
    return "Medium"


def _infer_method_path(text: str, slug: str) -> tuple[str, str]:
    t = text.lower()
    if any(w in t for w in ("create", "add", "submit", "register", "upload")):
        return "post", f"/{slug}s"
    if any(w in t for w in ("update", "edit", "modify", "change", "patch")):
        return "patch", f"/{slug}s/{{id}}"
    if any(w in t for w in ("delete", "remove", "cancel")):
        return "delete", f"/{slug}s/{{id}}"
    if any(w in t for w in ("list", "search", "filter", "get all")):
        return "get", f"/{slug}s"
    return "get", f"/{slug}s/{{id}}"


def _extract_sheet_rows(rows: list, mapping: dict, filename: str) -> list:
    us_col    = mapping.get("userStory", "")
    id_col    = mapping.get("storyId", "")
    title_col = mapping.get("title", "")
    pri_col   = mapping.get("priority", "")
    ac_col    = mapping.get("acceptanceCriteria", "")
    results = []
    for i, row in enumerate(rows):
        if not isinstance(row, dict):
            continue
        user_story = str(row.get(us_col, "")).strip() if us_col else ""
        if not user_story:
            continue
        story_id  = (str(row.get(id_col, "")).strip() if id_col else "") or f"FR-{i+1:03d}"
        title_val = str(row.get(title_col, "")).strip() if title_col else ""
        priority  = str(row.get(pri_col, "")).strip()  if pri_col  else ""
        criteria  = str(row.get(ac_col, "")).strip()   if ac_col   else ""
        slug      = re.sub(r"[^a-z0-9]+", "-", story_id.lower()).strip("-") or "resource"
        method, path = _infer_method_path(user_story, slug)
        ac_clean  = " ".join(c.strip() for c in re.split(r"[;\n]+", criteria) if c.strip())
        results.append({
            "id": story_id, "title": title_val or story_id,
            "desc": user_story, "source": f"Spreadsheet: {filename}",
            "priority": _priority_map(priority), "status": "Draft",
            "method": method, "path": path,
            "summary": user_story[:120], "acceptanceCriteria": ac_clean,
        })
    return results


async def _groq_sheet_fallback(rows: list, filename: str) -> list:
    text   = "\n".join(json.dumps(r) for r in rows[:80])
    prompt = (
        f'You are a business analyst. Extract structured API functional requirements '
        f'from these spreadsheet rows.\nReturn a JSON array only (no markdown) where each item has:\n'
        f'id, title, desc, source ("Spreadsheet: {filename}"), priority (High/Medium/Low), '
        f'status ("Draft"), method, path, summary, acceptanceCriteria (string).\nRows:\n{text[:6000]}'
    )
    raw     = await _get_groq()._call_groq(prompt)
    results = json.loads(_clean_json(raw))
    for r in results:
        if "desc" not in r and "description" in r:
            r["desc"] = r.pop("description")
        r.setdefault("desc", "")
        r.setdefault("summary", "")
        r.setdefault("method", "get")
        r.setdefault("path", "/resource")
        r.setdefault("acceptanceCriteria", "")
        r.setdefault("status", "Draft")
    return results


async def _extract_word(contents: bytes) -> tuple[list, str]:
    try:
        import docx
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx is not installed.")
    try:
        with io.BytesIO(contents) as buf:
            doc   = docx.Document(buf)
            parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_text:
                        parts.append(row_text)
            text = "\n".join(parts)
    except Exception as ex:
        raise HTTPException(status_code=400, detail=f"Failed to read document: {ex}")
    if not text.strip():
        raise HTTPException(status_code=400, detail="Document appears to be empty.")
    prompt = (
        'You are a business analyst. Extract structured functional requirements from this document.\n'
        'Return a JSON array only (no markdown) where each item has:\n'
        'id (FR-001…), title, desc (full user story), source ("Uploaded Document"),\n'
        'priority (High/Medium/Low), status ("Draft"), method, path, summary,\n'
        'acceptanceCriteria (array of strings — populate if criteria exist).\n\n'
        f'Document:\n{text[:8000]}'
    )
    try:
        raw          = await _get_groq()._call_groq(prompt)
        requirements = json.loads(_clean_json(raw))
        for r in requirements:
            if "desc" not in r and "description" in r:
                r["desc"] = r.pop("description")
            r.setdefault("desc", "")
            r.setdefault("summary", "")
            r.setdefault("method", "get")
            r.setdefault("path", "/resource")
            r.setdefault("status", "Draft")
            ac = r.get("acceptanceCriteria", [])
            r["acceptanceCriteria"] = (
                " ".join(str(c).strip() for c in ac if str(c).strip())
                if isinstance(ac, list) else (ac if isinstance(ac, str) else "")
            )
        return requirements, text
    except Exception as ex:
        logger.warning("LLM word extraction failed (%s), falling back to rule-based", ex)
        return _rule_based_extract_doc(text), text


def _rule_based_extract_doc(text: str) -> list:
    METHOD_KW = {
        "create": "post", "add": "post", "register": "post", "submit": "post", "upload": "post",
        "update": "put",  "edit": "put",  "modify": "patch", "change": "patch",
        "delete": "delete", "remove": "delete",
        "get": "get", "list": "get", "view": "get", "retrieve": "get", "search": "get",
    }
    PRIORITY_KW = {
        "critical": "High", "must": "High", "required": "High",
        "should": "Medium", "recommended": "Medium",
        "optional": "Low", "could": "Low",
    }
    lines     = [l.strip() for l in re.split(r"[\n•\-–]", text) if len(l.strip()) > 20]
    sentences, seen, results, counter = [], set(), [], 1
    for line in lines:
        for s in re.split(r"(?<=[.!?])\s+", line):
            if len(s.strip()) > 20:
                sentences.append(s.strip())
    for sentence in sentences:
        if sentence in seen or len(sentence) < 25 or sentence.isupper():
            continue
        seen.add(sentence)
        low      = sentence.lower()
        method   = next((m for kw, m in METHOD_KW.items()   if kw in low), "get")
        priority = next((p for kw, p in PRIORITY_KW.items() if kw in low), "Medium")
        words    = re.findall(r"[a-zA-Z]+", sentence)
        nouns    = [w.lower() for w in words if len(w) > 3 and w.lower() not in
                    {"shall", "should", "must", "will", "have", "with", "that",
                     "this", "from", "into", "able", "user", "users", "system", "allow"}]
        resource = nouns[0] if nouns else "resource"
        path     = f"/{resource}s" if method in ("get", "post") else f"/{resource}s/{{id}}"
        results.append({
            "id": f"FR-{counter:03d}", "title": " ".join(words[:6]).title(),
            "desc": sentence, "source": "Uploaded Document",
            "priority": priority, "status": "Draft",
            "method": method, "path": path,
            "summary": f"{method.upper()} {path}", "acceptanceCriteria": "",
        })
        counter += 1
        if counter > 50:
            break
    return results


@router.post("/extract")
async def extract_requirements(
    file: UploadFile = File(..., description=".docx, .xlsx, or .csv file"),
    userStory: Optional[str] = Form(None, description="Column name containing user story (required for Excel/CSV only)"),
):
    """
    Unified extraction endpoint for Word documents and Excel/CSV spreadsheets.
    Word (.docx) — requirements extracted by Groq AI immediately.
    Excel/CSV — provide userStory column name to extract; omit to trigger column-mapping modal.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file uploaded.")
    ext = ("." + file.filename.rsplit(".", 1)[-1]).lower() if "." in file.filename else ""
    if ext not in _ALL_EXTS:
        raise HTTPException(status_code=400, detail=f"Unsupported file type '{ext}'. Allowed: .docx, .xlsx, .csv")
    contents = await file.read()
    if not contents:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    # Build columnMapping JSON from the userStory field
    columnMapping = json.dumps({"userStory": userStory.strip()}) if userStory and userStory.strip() else None

    if ext in _WORD_EXTS:
        requirements, raw_text = await _extract_word(contents)
        return {"requirements": requirements, "raw_text": raw_text, "extraction_method": "word"}

    rows = _parse_sheet(contents, file.filename)
    if not rows:
        raise HTTPException(status_code=400, detail="No data rows found in the file.")
    columns = list(rows[0].keys())

    if not columnMapping or not columnMapping.strip():
        return {"needs_mapping": True, "columns": columns, "filename": file.filename}

    mapping = _build_mapping(columnMapping, columns)
    if not mapping.get("userStory"):
        raise HTTPException(
            status_code=400,
            detail=f"Could not find a user story column in: {columns}. Pass the exact column header name in columnMapping."
        )

    requirements = _extract_sheet_rows(rows, mapping, file.filename)
    if not requirements:
        try:
            requirements = await _groq_sheet_fallback(rows, file.filename)
        except Exception as ex:
            raise HTTPException(status_code=500, detail=f"Extraction failed: {ex}")
    if not requirements:
        raise HTTPException(status_code=400, detail="No user stories found. Check your column mapping.")

    return {
        "requirements": requirements, "total": len(requirements),
        "filename": file.filename, "extraction_method": "spreadsheet",
        "columns_detected": columns,
        "mapping_used": {k: v for k, v in mapping.items() if v},
    }


@router.post("/swagger-docs")
async def generate_swagger_docs(request: ArtifactRequest):
    if not request.open_api_yaml or not request.open_api_yaml.strip():
        raise HTTPException(status_code=400, detail="OpenApiYaml is required.")
    try:
        import yaml as pyyaml
        doc = pyyaml.safe_load(_clean_yaml(request.open_api_yaml))
        openapi_json = json.dumps(doc)
        title = doc.get("info", {}).get("title", "API Documentation")

        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
  <title>{title} - Swagger UI</title>
  <link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/swagger-ui-dist@5.11.0/swagger-ui.css" />
  <style>
    body {{ margin: 0; padding: 0; font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif; }}
    *, *::before, *::after {{ box-sizing: border-box; }}
    .swagger-ui .topbar {{ background: linear-gradient(90deg, #4f46e5, #6d28d9); border-bottom: 2px solid #3730a3; }}
    .swagger-ui .topbar .download-url-wrapper {{ display: none; }}
    .swagger-ui .info .title {{ color: #0f172a; font-size: 32px; }}
    .swagger-ui .opblock {{ border: 1px solid #e2e8f0; border-radius: 8px; margin-bottom: 12px; }}
    .swagger-ui .btn.execute {{ background: #4f46e5; color: #fff; border: none; font-weight: 600; padding: 8px 20px; border-radius: 6px; }}
  </style>
</head>
<body>
  <div id="swagger-ui"></div>
  <script src="https://cdn.jsdelivr.net/npm/swagger-ui-dist@5.11.0/swagger-ui-bundle.js"></script>
  <script src="https://cdn.jsdelivr.net/npm/swagger-ui-dist@5.11.0/swagger-ui-standalone-preset.js"></script>
  <script>
    window.onload = function() {{
      SwaggerUIBundle({{
        spec: {openapi_json},
        dom_id: '#swagger-ui',
        presets: [SwaggerUIBundle.presets.apis, SwaggerUIStandalonePreset],
        layout: 'BaseLayout',
        deepLinking: true,
        tryItOutEnabled: true,
        persistAuthorization: true,
      }});
    }};
  </script>
</body>
</html>"""
        return {"content": html, "file_name": "swagger_docs.html", "content_type": "text/html"}
    except Exception as ex:
        logger.error("Swagger docs generation failed: %s", ex)
        raise HTTPException(status_code=500, detail=f"Swagger docs generation failed: {ex}")


@router.post("/data-models")
async def generate_data_models(request: ArtifactRequest):
    if not request.open_api_yaml or not request.open_api_yaml.strip():
        raise HTTPException(status_code=400, detail="OpenApiYaml is required.")

    prompt = f"""You are an API architect. Extract all data models and schemas from the following OpenAPI specification.

Return a single JSON object only (no markdown, no explanation) where:
- Each key is a schema/model name
- Each value is an object with:
  - description: one-sentence description of the model
  - type: the JSON schema type (object, array, string, etc.)
  - properties: object where each key is a field name with value containing: type, description, required (bool), example
  - required: array of required field names

OpenAPI YAML:
{request.open_api_yaml[:8000]}
"""
    try:
        raw = await _get_groq()._call_groq(prompt)
        cleaned = _clean_json(raw)
        models = json.loads(cleaned)
        return {"content": json.dumps(models, indent=2), "file_name": "data_models.json", "content_type": "application/json"}
    except Exception as ex:
        logger.error("Data model generation failed: %s", ex)
        raise HTTPException(status_code=500, detail=f"Data model generation failed: {ex}")
