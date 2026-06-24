"""
Unified document extraction endpoint.
Handles .docx (Word) and .xlsx/.csv (Excel/CSV) in a single route.
"""
import csv
import io
import json
import logging
import re
from typing import Optional

import openpyxl
from fastapi import APIRouter, File, Form, HTTPException, UploadFile

from dependencies import get_groq_service
from routers.designer import _clean_json

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/designer", tags=["documents"])

WORD_EXTS  = {".docx"}
SHEET_EXTS = {".xlsx", ".csv"}
ALL_EXTS   = WORD_EXTS | SHEET_EXTS

# Patterns used to auto-detect each logical column from raw header names
_COL_PATTERNS: dict[str, list[str]] = {
    "userStory":          ["user story", "userstory", "user_story", "description",
                           "desc", "story", "requirement", "as a", "as an"],
    "storyId":            ["story id", "storyid", "story_id", "id", "ticket", "issue"],
    "title":              ["title", "name", "summary", "epic"],
    "priority":           ["priority", "severity", "importance"],
    "acceptanceCriteria": ["acceptance criteria", "acceptance_criteria",
                           "acceptancecriteria", "ac", "criteria", "done when",
                           "definition of done", "dod"],
}


def _get_groq():
    return get_groq_service()


# ── sheet parsing ─────────────────────────────────────────────────────────────

def _parse_sheet(content: bytes, filename: str) -> list[dict]:
    ext = ("." + filename.rsplit(".", 1)[-1]).lower() if "." in filename else ""
    if ext == ".csv":
        text = content.decode("utf-8-sig", errors="replace")
        return [dict(row) for row in csv.DictReader(io.StringIO(text))]
    if ext == ".xlsx":
        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        ws = wb.active
        rows = list(ws.iter_rows(values_only=True))
        wb.close()
        if not rows:
            return []
        headers = [str(h).strip() if h is not None else f"col_{i}"
                   for i, h in enumerate(rows[0])]
        return [
            {headers[j]: (str(cell).strip() if cell is not None else "")
             for j, cell in enumerate(row)}
            for row in rows[1:]
        ]
    raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}")


# ── column resolution ─────────────────────────────────────────────────────────

def _resolve_column(value: str, columns: list[str]) -> str:
    """Match a user-supplied value to a real column header.
    Priority: exact → case-insensitive exact → partial containment → not found.
    """
    if not value:
        return ""
    if value in columns:
        return value
    vl = value.strip().lower()
    for col in columns:
        if col.strip().lower() == vl:
            return col
    for col in columns:
        cl = col.strip().lower()
        if vl in cl or cl in vl:
            return col
    return ""


def _auto_detect(columns: list[str]) -> dict[str, str]:
    """Auto-detect the best column header for every logical field."""
    result: dict[str, str] = {}
    for field, patterns in _COL_PATTERNS.items():
        for col in columns:
            cl = col.strip().lower()
            if any(p in cl for p in patterns):
                result[field] = col
                break
    return result


def _build_mapping(raw_input: str, columns: list[str]) -> dict[str, str]:
    """
    Convert whatever the caller sends into a resolved column mapping dict.

    Accepted input formats:
      • JSON object  – {"userStory":"User Story","acceptanceCriteria":"AC", ...}
      • Plain string – "User Story"  (treated as the userStory column name)
      • Keyword      – "userStory"   (auto-detect from headers)

    In ALL cases the optional fields (storyId, title, priority, acceptanceCriteria)
    are filled in via auto-detection if not explicitly provided.
    """
    raw = raw_input.strip()

    if raw.startswith("{"):
        # ── JSON object ──────────────────────────────────────────────────────
        try:
            supplied: dict = json.loads(raw)
        except Exception:
            raise HTTPException(status_code=400,
                                detail="columnMapping must be valid JSON when sent as an object.")
        mapping = {k: _resolve_column(v, columns) for k, v in supplied.items() if v}
    else:
        # ── Plain string or camelCase keyword ────────────────────────────────
        # Try to resolve it as a literal column name first
        resolved = _resolve_column(raw, columns)
        mapping = {"userStory": resolved} if resolved else {}

    # Fill every missing optional field via auto-detection
    auto = _auto_detect(columns)
    for field in _COL_PATTERNS:
        if not mapping.get(field):
            mapping[field] = auto.get(field, "")

    return mapping


# ── extraction helpers ────────────────────────────────────────────────────────

def _priority_map(raw: str) -> str:
    r = raw.strip().lower()
    if r in ("high", "1", "critical", "blocker"):
        return "High"
    if r in ("low", "3", "4", "minor", "trivial"):
        return "Low"
    return "Medium"


def _infer_method_path(text: str, slug: str) -> tuple[str, str]:
    t = text.lower()
    if any(w in t for w in ("create", "add", "submit", "register", "upload")):
        return "post", f"/{slug}s"
    if any(w in t for w in ("update", "edit", "modify", "change", "patch")):
        return "patch", f"/{slug}s/{{id}}"
    if any(w in t for w in ("delete", "remove", "cancel")):
        return "delete", f"/{slug}s/{{id}}"
    if any(w in t for w in ("list", "search", "filter", "get all")):
        return "get", f"/{slug}s"
    return "get", f"/{slug}s/{{id}}"


def _extract_sheet_rows(rows: list, mapping: dict, filename: str) -> list:
    us_col    = mapping.get("userStory", "")
    id_col    = mapping.get("storyId", "")
    title_col = mapping.get("title", "")
    pri_col   = mapping.get("priority", "")
    ac_col    = mapping.get("acceptanceCriteria", "")

    results = []
    for i, row in enumerate(rows):
        if not isinstance(row, dict):
            continue

        user_story = str(row.get(us_col, "")).strip() if us_col else ""
        if not user_story:
            continue

        story_id = (str(row.get(id_col, "")).strip() if id_col else "") or f"FR-{i+1:03d}"
        title_val = str(row.get(title_col, "")).strip() if title_col else ""
        priority  = str(row.get(pri_col, "")).strip()  if pri_col  else ""
        criteria  = str(row.get(ac_col, "")).strip()   if ac_col   else ""

        slug = re.sub(r"[^a-z0-9]+", "-", story_id.lower()).strip("-") or "resource"
        method, path = _infer_method_path(user_story, slug)

        # Keep acceptanceCriteria as its own field (not merged into desc)
        ac_clean = " ".join(
            c.strip() for c in re.split(r"[;\n]+", criteria) if c.strip()
        )

        results.append({
            "id":                 story_id,
            "title":              title_val or story_id,
            "desc":               user_story,
            "source":             f"Spreadsheet: {filename}",
            "priority":           _priority_map(priority),
            "status":             "Draft",
            "method":             method,
            "path":               path,
            "summary":            user_story[:120],
            "acceptanceCriteria": ac_clean,
        })
    return results


async def _groq_sheet_fallback(rows: list, filename: str) -> list:
    text   = "\n".join(json.dumps(r) for r in rows[:80])
    prompt = (
        f'You are a business analyst. Extract structured API functional requirements '
        f'from these spreadsheet rows.\n'
        f'Return a JSON array only (no markdown) where each item has:\n'
        f'id, title, desc, source ("Spreadsheet: {filename}"), '
        f'priority (High/Medium/Low), status ("Draft"), '
        f'method, path, summary, acceptanceCriteria (string).\n'
        f'Rows:\n{text[:6000]}'
    )
    raw     = await _get_groq()._call_groq(prompt)
    results = json.loads(_clean_json(raw))
    for r in results:
        if "desc" not in r and "description" in r:
            r["desc"] = r.pop("description")
        r.setdefault("desc", "")
        r.setdefault("summary", "")
        r.setdefault("method", "get")
        r.setdefault("path", "/resource")
        r.setdefault("acceptanceCriteria", "")
        r.setdefault("status", "Draft")
    return results


async def _extract_word(contents: bytes) -> tuple[list, str]:
    try:
        import docx
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx is not installed.")

    try:
        with io.BytesIO(contents) as buf:
            doc   = docx.Document(buf)
            parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        c.text.strip() for c in row.cells if c.text.strip()
                    )
                    if row_text:
                        parts.append(row_text)
            text = "\n".join(parts)
    except Exception as ex:
        raise HTTPException(status_code=400, detail=f"Failed to read document: {ex}")

    if not text.strip():
        raise HTTPException(status_code=400, detail="Document appears to be empty.")

    prompt = (
        'You are a business analyst. Extract structured functional requirements '
        'from this document.\n'
        'Return a JSON array only (no markdown) where each item has:\n'
        'id (FR-001…), title, desc (full user story), source ("Uploaded Document"),\n'
        'priority (High/Medium/Low), status ("Draft"), method, path, summary,\n'
        'acceptanceCriteria (array of strings — populate if criteria exist).\n\n'
        f'Document:\n{text[:8000]}'
    )
    try:
        raw          = await _get_groq()._call_groq(prompt)
        requirements = json.loads(_clean_json(raw))
        for r in requirements:
            if "desc" not in r and "description" in r:
                r["desc"] = r.pop("description")
            r.setdefault("desc", "")
            r.setdefault("summary", "")
            r.setdefault("method", "get")
            r.setdefault("path", "/resource")
            r.setdefault("status", "Draft")
            ac = r.get("acceptanceCriteria", [])
            r["acceptanceCriteria"] = (
                " ".join(str(c).strip() for c in ac if str(c).strip())
                if isinstance(ac, list) else (ac if isinstance(ac, str) else "")
            )
        return requirements, text
    except Exception as ex:
        logger.warning("LLM word extraction failed (%s), falling back to rule-based", ex)
        return _rule_based_extract(text), text


def _rule_based_extract(text: str) -> list:
    METHOD_KW = {
        "create": "post", "add": "post", "register": "post",
        "submit": "post", "upload": "post",
        "update": "put",  "edit": "put",  "modify": "patch", "change": "patch",
        "delete": "delete", "remove": "delete",
        "get": "get", "list": "get", "view": "get", "retrieve": "get", "search": "get",
    }
    PRIORITY_KW = {
        "critical": "High", "must": "High", "required": "High",
        "should": "Medium", "recommended": "Medium",
        "optional": "Low", "could": "Low",
    }
    lines = [l.strip() for l in re.split(r"[\n•\-–]", text) if len(l.strip()) > 20]
    sentences, seen, results, counter = [], set(), [], 1
    for line in lines:
        for s in re.split(r"(?<=[.!?])\s+", line):
            if len(s.strip()) > 20:
                sentences.append(s.strip())
    for sentence in sentences:
        if sentence in seen or len(sentence) < 25 or sentence.isupper():
            continue
        seen.add(sentence)
        low      = sentence.lower()
        method   = next((m for kw, m in METHOD_KW.items()   if kw in low), "get")
        priority = next((p for kw, p in PRIORITY_KW.items() if kw in low), "Medium")
        words    = re.findall(r"[a-zA-Z]+", sentence)
        nouns    = [w.lower() for w in words if len(w) > 3 and w.lower() not in
                    {"shall", "should", "must", "will", "have", "with", "that",
                     "this", "from", "into", "able", "user", "users", "system", "allow"}]
        resource = nouns[0] if nouns else "resource"
        path     = f"/{resource}s" if method in ("get", "post") else f"/{resource}s/{{id}}"
        results.append({
            "id": f"FR-{counter:03d}", "title": " ".join(words[:6]).title(),
            "desc": sentence, "source": "Uploaded Document",
            "priority": priority, "status": "Draft",
            "method": method, "path": path,
            "summary": f"{method.upper()} {path}", "acceptanceCriteria": "",
        })
        counter += 1
        if counter > 50:
            break
    return results


# ── endpoint ──────────────────────────────────────────────────────────────────

@router.post("/extract")
async def extract_requirements(
    file: UploadFile = File(..., description=".docx, .xlsx, or .csv file"),
    columnMapping: Optional[str] = Form(
        None,
        description=(
            "Column mapping for Excel/CSV. Accepted formats:\n"
            "1. Omit entirely  → returns {needs_mapping, columns} so UI shows mapping modal.\n"
            "2. Plain string   → 'User Story'  or keyword 'userStory'  "
            "(auto-detects all other columns).\n"
            "3. JSON object    → {\"userStory\":\"User Story\",\"acceptanceCriteria\":\"AC\","
            "\"storyId\":\"ID\",\"title\":\"Title\",\"priority\":\"Priority\"}.\n"
            "Only userStory is required; all other fields are auto-detected if omitted."
        ),
    ),
):
    """
    Unified extraction endpoint for Word documents and Excel/CSV spreadsheets.

    **Word (.docx)** — file uploaded, requirements extracted by Groq AI immediately.

    **Excel/CSV (.xlsx/.csv)**
    - No `columnMapping` → `{needs_mapping: true, columns: [...]}` (show mapping modal)
    - With `columnMapping` → extracts requirements; missing optional columns are auto-detected
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file uploaded.")

    filename = file.filename
    ext = ("." + filename.rsplit(".", 1)[-1]).lower() if "." in filename else ""
    if ext not in ALL_EXTS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. Allowed: .docx, .xlsx, .csv"
        )

    contents = await file.read()
    if not contents:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    # ── Word ──────────────────────────────────────────────────────────────────
    if ext in WORD_EXTS:
        requirements, raw_text = await _extract_word(contents)
        return {"requirements": requirements, "raw_text": raw_text,
                "extraction_method": "word"}

    # ── Spreadsheet ───────────────────────────────────────────────────────────
    rows = _parse_sheet(contents, filename)
    if not rows:
        raise HTTPException(status_code=400, detail="No data rows found in the file.")

    columns = list(rows[0].keys())

    # No mapping at all → tell the frontend to show the column-mapping modal
    if not columnMapping or not columnMapping.strip():
        return {"needs_mapping": True, "columns": columns, "filename": filename}

    # Build the resolved mapping (handles plain string, keyword, JSON object)
    mapping = _build_mapping(columnMapping, columns)

    if not mapping.get("userStory"):
        raise HTTPException(
            status_code=400,
            detail=(
                f"Could not find a user story column in: {columns}. "
                "Pass the exact column header name (e.g. 'User Story') in columnMapping."
            )
        )

    logger.info(
        "Extracting %s | mapping=%s", filename,
        {k: v for k, v in mapping.items() if v}
    )

    requirements = _extract_sheet_rows(rows, mapping, filename)

    if not requirements:
        try:
            requirements = await _groq_sheet_fallback(rows, filename)
        except Exception as ex:
            logger.error("Groq fallback failed: %s", ex)
            raise HTTPException(status_code=500, detail=f"Extraction failed: {ex}")

    if not requirements:
        raise HTTPException(
            status_code=400,
            detail="No user stories found. Check your column mapping."
        )

    return {
        "requirements":      requirements,
        "total":             len(requirements),
        "filename":          filename,
        "extraction_method": "spreadsheet",
        "columns_detected":  columns,
        "mapping_used":      {k: v for k, v in mapping.items() if v},
    }
